from docx import Document

def extract_docx_content(file_path):
    doc = Document(file_path)
    for para in doc.paragraphs:
        print(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' '.join(cell.text for cell in row.cells)
            print(row_text)

if __name__ == "__main__":
    extract_docx_content("my_competence.docx")
